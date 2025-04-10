import os
import subprocess
import uuid
import tempfile
import logging
from docx import Document
import fitz  # PyMuPDF

def word_to_pdf(doc_path):
    """
    Convert a Word document to PDF.
    Returns the path to the created PDF file.
    """
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Word document not found: {doc_path}")
    
    output_dir = os.path.dirname(doc_path)
    output_filename = f"{uuid.uuid4()}.pdf"
    output_path = os.path.join(output_dir, output_filename)
    
    # Try using LibreOffice if available (common in Linux environments)
    try:
        # Create a temporary directory for LibreOffice conversion
        temp_dir = tempfile.mkdtemp()
        
        # Run LibreOffice headless conversion
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf",
            "--outdir", temp_dir, doc_path
        ], check=True, stderr=subprocess.PIPE)
        
        # LibreOffice keeps the original filename, just changes extension
        base_name = os.path.basename(doc_path)
        name_without_ext = os.path.splitext(base_name)[0]
        libreoffice_output = os.path.join(temp_dir, f"{name_without_ext}.pdf")
        
        # Copy to our output location with UUID filename
        if os.path.exists(libreoffice_output):
            with open(libreoffice_output, 'rb') as src, open(output_path, 'wb') as dst:
                dst.write(src.read())
            return output_path
    except Exception as e:
        logging.warning(f"LibreOffice conversion failed: {str(e)}")
    
    # Fallback method using python-docx and reportlab
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        # Extract text from the Word document
        doc = Document(doc_path)
        
        # Create a PDF with the extracted text
        pdf_document = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create a basic style for document text
        normal_style = styles["Normal"]
        title_style = styles["Heading1"]
        
        # Elements to add to PDF
        elements = []
        
        # Add document content
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                # Use heading style for headings
                elements.append(Paragraph(para.text, title_style))
                elements.append(Spacer(1, 12))
            elif para.text.strip():  # Only add non-empty paragraphs
                elements.append(Paragraph(para.text, normal_style))
                elements.append(Spacer(1, 6))
        
        # Build the PDF
        pdf_document.build(elements)
        
        return output_path
    except Exception as e:
        logging.error(f"Fallback conversion failed: {str(e)}")
        raise

def pdf_to_word(pdf_path):
    """
    Extract text from a PDF and create a Word document.
    Returns the path to the created Word document.
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")
    
    output_dir = os.path.dirname(pdf_path)
    output_filename = f"{uuid.uuid4()}.docx"
    output_path = os.path.join(output_dir, output_filename)
    
    try:
        # Open the PDF with PyMuPDF
        doc = fitz.open(pdf_path)
        
        # Create a new Word document
        word_doc = Document()
        
        # Add a title to the document
        original_filename = os.path.basename(pdf_path)
        word_doc.add_heading(f"Converted from {original_filename}", 0)
        
        # Extract text from each page and add to the Word document
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            
            # Add page number header
            word_doc.add_heading(f"Page {page_num + 1}", level=1)
            
            # Get text content
            try:
                # First try to get structured blocks
                blocks = page.get_text("blocks")
                
                # Process each text block
                for block in blocks:
                    if len(block) >= 5:  # Make sure the block has enough elements
                        text = block[4]  # The text content is at index 4
                        if text and text.strip():  # Skip empty blocks
                            # Add the text to the Word document
                            word_doc.add_paragraph(text.strip())
                    
            except Exception as block_error:
                logging.warning(f"Error getting text blocks: {str(block_error)}")
                # Fallback to simple text extraction
                text = page.get_text()
                for paragraph in text.split('\n\n'):
                    if paragraph.strip():
                        word_doc.add_paragraph(paragraph.strip())
            
            # Add a page break after each page (except the last)
            if page_num < len(doc) - 1:
                word_doc.add_page_break()
        
        # Save the Word document
        word_doc.save(output_path)
        
        return output_path
    except Exception as e:
        logging.error(f"Error converting PDF to Word: {str(e)}")
        raise
